"""Document upload/ingest, filesystem browse/import/scan, vision, and settings endpoints."""

from __future__ import annotations

import asyncio
import hashlib
import io
import logging
import re
import shutil
import time
from pathlib import Path
from typing import Optional

import httpx
from fastapi import APIRouter, HTTPException, Query, UploadFile, File, Form
from fastapi.responses import JSONResponse

from core.config import settings

from core.artifacts.document import Document, DocumentChunk, DocumentType, DocumentStatus
from core.artifacts.storage_documents_pg import (
    save_document,
    load_document,
    load_all_documents,
    save_chunk,
    load_chunks_by_document,
    delete_document,
    find_document_by_hash,
)
from core.tools.mcp_filesystem import FilesystemAccess
from core.tools.document_ingest import DocumentIngestor

log = logging.getLogger("noesis.documents")

router = APIRouter(tags=["documents"])


# ---------------------------------------------------------------------------
# Document upload & management
# ---------------------------------------------------------------------------

@router.post("/documents/upload")
async def upload_document(
    file: UploadFile = File(...),
    tags: str = Form(default=""),
    notes: str = Form(default=""),
):
    """Upload document for indexing."""
    content = await file.read()

    if len(content) > 50 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="File too large (max 50MB)")

    file_hash = hashlib.sha256(content).hexdigest()

    existing = await find_document_by_hash(file_hash)
    if existing:
        return {"status": "duplicate", "document_id": existing.id}

    ext = Path(file.filename).suffix[1:].lower()
    try:
        file_type = DocumentType(ext)
    except ValueError:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type: .{ext}. Supported: PDF, DOCX, TXT, MD, CSV, images, SVG"
        )

    doc = Document(
        filename=file.filename,
        filepath=f"uploads/{file.filename}",
        file_type=file_type,
        file_size=len(content),
        file_hash=file_hash,
        tags=[t.strip() for t in tags.split(",") if t.strip()],
        notes=notes,
    )
    await save_document(doc)

    upload_dir = Path("data/documents/uploads") / doc.id
    upload_dir.mkdir(parents=True, exist_ok=True)
    (upload_dir / file.filename).write_bytes(content)

    asyncio.create_task(_index_document_task(doc.id, content))

    return {"status": "uploaded", "document_id": doc.id}


async def _index_document_task(doc_id: str, content: bytes):
    """Background task: ingest document into RAG system."""
    from core.rag.vector_store import PgVectorStore

    try:
        doc = await load_document(doc_id)
        if not doc:
            log.error("Document %s not found for indexing", doc_id)
            return

        ingestor = DocumentIngestor(PgVectorStore())
        await ingestor.ingest(doc, content)
        log.info("Indexed %s (%d chunks)", doc.filename, doc.chunk_count)
    except Exception as e:
        log.error("Failed to index document %s: %s", doc_id, e, exc_info=True)


@router.get("/documents")
async def list_documents(
    status: Optional[str] = Query(default=None),
    limit: int = Query(default=100, ge=1, le=500),
):
    """List documents with optional status filter."""
    doc_status = DocumentStatus(status) if status else None
    docs = await load_all_documents(status=doc_status, limit=limit)
    return [doc.to_dict() for doc in docs]


@router.get("/documents/{doc_id}")
async def get_document(doc_id: str):
    """Get document metadata by ID."""
    doc = await load_document(doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    return doc.to_dict()


@router.get("/documents/{doc_id}/chunks")
async def get_document_chunks(
    doc_id: str,
    limit: Optional[int] = Query(default=None, ge=1, le=1000),
):
    """Get all chunks for a document."""
    doc = await load_document(doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    chunks = await load_chunks_by_document(doc_id, limit=limit)
    return [chunk.to_dict() for chunk in chunks]


@router.delete("/documents/{doc_id}")
async def remove_document(doc_id: str):
    """Delete document and all its chunks."""
    from core.rag.vector_store import PgVectorStore

    doc = await load_document(doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    # Remove chunks from vector store
    store = PgVectorStore()
    chunks = await load_chunks_by_document(doc_id)
    for chunk in chunks:
        await store.remove(f"chunk:{chunk.id}")

    # Delete from PostgreSQL (CASCADE handles chunks)
    await delete_document(doc_id)

    # Delete uploaded file from server disk
    upload_dir = Path("data/documents/uploads") / doc_id
    if upload_dir.exists():
        shutil.rmtree(upload_dir)

    return {"status": "deleted", "document_id": doc_id}


# ---------------------------------------------------------------------------
# Filesystem browse/import/scan
# ---------------------------------------------------------------------------

@router.post("/documents/browse")
async def browse_filesystem(req: dict):
    """Browse allowed directories on server filesystem."""
    fs = FilesystemAccess()

    directory = req.get("directory", str(Path.home() / "Documents"))
    recursive = req.get("recursive", False)

    try:
        files = fs.list_files(directory, recursive)
        return {"files": files}
    except PermissionError as e:
        raise HTTPException(status_code=403, detail=str(e))
    except FileNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e))


@router.post("/documents/import")
async def import_from_filesystem(req: dict):
    """Import file from server filesystem path."""
    fs = FilesystemAccess()
    filepath = req.get("filepath")

    if not filepath:
        raise HTTPException(status_code=400, detail="filepath required")

    try:
        content = fs.read_file(filepath)
        file_hash = fs.compute_hash(content)

        existing = await find_document_by_hash(file_hash)
        if existing:
            return {"status": "duplicate", "document_id": existing.id}

        ext = Path(filepath).suffix[1:].lower()
        try:
            file_type = DocumentType(ext)
        except ValueError:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type: .{ext}"
            )

        doc = Document(
            filename=Path(filepath).name,
            filepath=filepath,
            file_type=file_type,
            file_size=len(content),
            file_hash=file_hash,
        )
        await save_document(doc)

        asyncio.create_task(_index_document_task(doc.id, content))

        return {"status": "imported", "document_id": doc.id}

    except PermissionError as e:
        raise HTTPException(status_code=403, detail=str(e))
    except FileNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@router.post("/documents/scan")
async def scan_and_import_directory(req: dict):
    """Scan a directory tree and import all supported files."""
    fs = FilesystemAccess()
    directory = req.get("directory")

    if not directory:
        raise HTTPException(status_code=400, detail="directory required")

    try:
        files = fs.list_files(directory, recursive=True)
    except PermissionError as e:
        raise HTTPException(status_code=403, detail=str(e))
    except FileNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e))

    results = []
    imported = 0
    duplicates = 0
    errors = 0

    for f in files:
        filepath = f["path"]
        try:
            content = fs.read_file(filepath)
            file_hash = fs.compute_hash(content)

            existing = await find_document_by_hash(file_hash)
            if existing:
                results.append({"path": filepath, "status": "duplicate", "document_id": existing.id})
                duplicates += 1
                continue

            ext = Path(filepath).suffix[1:].lower()
            try:
                file_type = DocumentType(ext)
            except ValueError:
                results.append({"path": filepath, "status": "error", "document_id": None})
                errors += 1
                continue

            doc = Document(
                filename=Path(filepath).name,
                filepath=filepath,
                file_type=file_type,
                file_size=len(content),
                file_hash=file_hash,
            )
            await save_document(doc)

            asyncio.create_task(_index_document_task(doc.id, content))

            results.append({"path": filepath, "status": "imported", "document_id": doc.id})
            imported += 1

        except Exception as e:
            results.append({"path": filepath, "status": "error", "document_id": None})
            errors += 1

    return {
        "scanned": len(files),
        "imported": imported,
        "duplicates": duplicates,
        "errors": errors,
        "files": results,
    }


# ---------------------------------------------------------------------------
# Vision endpoints
# ---------------------------------------------------------------------------

@router.post("/vision/caption")
async def caption_image(file: UploadFile = File(...)):
    """Generate a caption and OCR text for an uploaded image."""
    from core.vision.captioner import caption_and_ocr

    content = await file.read()
    if len(content) > 50 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="Image too large (max 50MB)")

    ext = Path(file.filename or "").suffix.lower()
    if ext not in (".jpg", ".jpeg", ".png"):
        raise HTTPException(status_code=400, detail="Unsupported image format. Use JPG or PNG.")

    try:
        result = caption_and_ocr(content)
        return {
            "caption": result["caption"],
            "ocr_text": result["ocr_text"],
            "filename": file.filename,
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Vision processing error: {str(e)}")


# ---------------------------------------------------------------------------
# FoxDigest compatibility — document summarization via LLM
# ---------------------------------------------------------------------------

def _extract_text_from_file(filename: str, content: bytes) -> str:
    """Extract text from an uploaded document for summarization."""
    ext = filename.lower()

    if ext.endswith(".docx"):
        from docx import Document as DocxDocument
        doc = DocxDocument(io.BytesIO(content))
        text = "\n".join(para.text for para in doc.paragraphs)
    elif ext.endswith(".pdf"):
        import fitz
        with fitz.open(stream=content, filetype="pdf") as pdf:
            text = "\n".join(page.get_text() for page in pdf)
    elif ext.endswith((".txt", ".md")):
        text = content.decode(errors="ignore")
    elif ext.endswith((".png", ".jpg", ".jpeg")):
        from PIL import Image
        import pytesseract
        image = Image.open(io.BytesIO(content))
        text = pytesseract.image_to_string(image, lang="eng")
    else:
        raise HTTPException(status_code=400, detail="Unsupported file format.")

    if len(text.strip()) < 20:
        raise HTTPException(status_code=422, detail="Unable to extract usable text from document.")

    return text.strip()


def _split_thinking(raw: str) -> str:
    """Strip <think> reasoning from DeepSeek-R1 output, return visible answer."""
    if "</think>" in raw:
        parts = raw.split("</think>", 1)
        return (parts[1].strip() if len(parts) > 1 else "")
    return raw.strip()


@router.post("/process-document")
async def process_document(file: UploadFile = File(...)):
    """Summarize an uploaded document using the local LLM.

    Compatible with FoxDigest's localLLM.js client.
    Accepts: PDF, DOCX, TXT, MD, PNG, JPG.
    Returns: {summary, key_points, processing_time, model_used}
    """
    try:
        start = time.time()
        content = await file.read()

        if len(content) > 10 * 1024 * 1024:
            raise HTTPException(status_code=413, detail="File too large (max 10MB)")

        extracted_text = _extract_text_from_file(file.filename, content)

        # Truncate to fit context window
        max_chars = 12000
        truncated_text = extracted_text[:max_chars]

        # Adjust summary length based on input
        if len(truncated_text) < 1000:
            length_instruction = "in under 75 words"
        elif len(truncated_text) < 3000:
            length_instruction = "in 100-150 words"
        else:
            length_instruction = "in 200-250 words"

        prompt = (
            f"You are a professional summarization assistant. Summarize the following document clearly "
            f"and concisely, {length_instruction}, using Markdown.\n\n"
            "Use this format:\n\n"
            "**Overview**\n[Brief description]\n\n"
            "**Details**\n[2-4 sentence explanation]\n\n"
            "**Outcome**\n[Optional: result, insight, or next steps.]\n\n"
            "### Key Points\n"
            "- [Point 1]\n- [Point 2]\n- [Point 3]\n\n"
            "Do not explain your response. Return only formatted markdown.\n\n"
            f"Document:\n{truncated_text}"
        )

        payload = {
            "model": settings.vllm_model_name,
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.3,
            "max_tokens": 1000,
        }

        async with httpx.AsyncClient(timeout=settings.llm_timeout) as client:
            resp = await client.post(
                f"{settings.vllm_base_url}/v1/chat/completions",
                json=payload,
            )
            resp.raise_for_status()
            data = resp.json()

        raw_answer = data["choices"][0]["message"]["content"].strip()
        visible_answer = _split_thinking(raw_answer)

        # Extract key points
        key_points_match = re.search(
            r"(?:#{1,3}\s*)?(?:\*\*)?Key Points(?:\*\*)?\s*\n((?:- .+\n?)+)",
            visible_answer,
        )
        key_points = []
        if key_points_match:
            key_points = [
                line.strip("- ").strip()
                for line in key_points_match.group(1).splitlines()
                if line.strip()
            ]

        elapsed = round(time.time() - start, 2)

        return JSONResponse(content={
            "summary": visible_answer,
            "key_points": key_points,
            "processing_time": elapsed,
            "model_used": "DeepSeek-R1-70B",
        })

    except HTTPException:
        raise
    except Exception as e:
        log.error("process-document failed: %s", e, exc_info=True)
        raise HTTPException(status_code=500, detail=f"Processing error: {str(e)}")


# ---------------------------------------------------------------------------
# Filesystem settings
# ---------------------------------------------------------------------------

@router.get("/settings/filesystem/allowed-roots")
async def get_allowed_roots():
    """Get list of allowed directory roots."""
    from core.settings.filesystem_settings import load_allowed_roots
    return {"roots": load_allowed_roots()}


@router.put("/settings/filesystem/allowed-roots")
async def update_allowed_roots(req: dict):
    """Update allowed directory roots."""
    from core.settings.filesystem_settings import save_allowed_roots

    roots = req.get("roots", [])
    if not isinstance(roots, list):
        raise HTTPException(status_code=400, detail="roots must be a list")

    save_allowed_roots(roots)
    return {"status": "updated"}
